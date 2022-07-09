from docx import *
from docx.enum.style import *
from docx.shared import *
from docx.oxml.ns import *
from docx.enum.section import WD_SECTION
from docx.enum.text import *

document = Document()
sections = document.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

section = document.sections[0]

sectPR =  section._sectPr
cols = sectPR.xpath('./w:cols')[0]

cols.set(qn('w:num'), '1')

document.add_heading('Balance Sheet', 0)

global content
content = 'Hello, thank you for using Py-accounting software. You May Delete this Message.'

def writedocx(content, font_name = 'Times New Roman', font_size = 12, font_bold = False, font_italic = False, font_underline = False, color = RGBColor(0, 0, 0),
              before_spacing = 5, after_spacing = 5, line_spacing = 1.5, keep_together = True, keep_with_next = False, page_break_before = False,
              widow_control = False, align = 'left', style = ''):
    paragraph = document.add_paragraph(str(content))
    paragraph.style = document.styles.add_style(style, WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = font_bold
    font.italic = font_italic
    font.underline = font_underline
    font.color.rgb = color
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(before_spacing)
    paragraph_format.space_after = Pt(after_spacing)
    paragraph.line_spacing = line_spacing
    paragraph_format.keep_together = keep_together
    paragraph_format.keep_with_next = keep_with_next
    paragraph_format.page_break_before = page_break_before
    paragraph_format.widow_control = widow_control
    if align.lower() == 'left':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif align.lower() == 'center':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif align.lower() == 'right':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    elif align.lower() == 'justify':
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    else:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
writedocx(content)

document.save('Word.docx')
