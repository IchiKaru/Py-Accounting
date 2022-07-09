#phase 1 of the project

import time
from docx import *
from docx.enum.style import *
from docx.shared import *
from docx.oxml.ns import *
from docx.enum.section import WD_SECTION
from docx.enum.text import *
time.sleep(0.5)

time.sleep(1)
def variables():
#global assets list
    global Current_Assets
    Current_Assets=[]
    global Current_Assets_Value
    Current_Assets_Value=[]
    global Fixed_Assets
    Fixed_Assets=[]
    global Fixed_Assets_Value
    Fixed_Assets_Value=[]
    global Intangible_Assets
    Intangible_Assets=[]
    global Intangible_Assets_Value
    Intangible_Assets_Value=[]
    #global liability list
    global Current_Liabilities
    Current_Liabilities=[]
    global Current_Liabilities_Amount
    Current_Liabilities_Amount=[]
    #global Capital_Value (or equity depending on your region) List
    global Capital_Entry
    Capital_Entry=[]
    global Capital_Value
    Capital_Value=[]
    #extended global variables
    global Revenues
    Revenues=[]
    global Revenue_Quantity
    Revenue_Quantity=[]
    global Expenditures
    Expenditures=[]
    global Expenditure_Cost
    Expenditure_Cost=[]
    global Withdrawal
    Withdrawal=[]
    global Withdrawal_value
    Withdrawal_value=[]
    #global variables
    global user_choice

def start_program(): #The very first thing that happens in the process. 
    print('What do you choose to Record?')
    time.sleep(1.5)
    user_choice = input("Enter 'Assets', 'Liabilities', or 'Capital':")
    print('-'*20)
    time.sleep(0.5)

    if user_choice == 'Assets': #If the user entered 'Assets' in the first question.
        asset_entry = input("What entry of assets is it? 'F' for fixed, 'C' for current, 'I' for intangible:")
        if asset_entry == 'C':
            current_asset_record()
        elif asset_entry == 'F':
            fixed_asset_record()
        elif asset_entry =='I':
            intangible_asset_record()
        else: #If neither of the user's choice was 'C', 'F', or 'I'.
            print('Sorry, that is an invalid syntax.')
            time.sleep(0.5)
            start_program()
    elif user_choice == 'Capital': #if the user entered 'Capital' in the first question.
        capital_entry = input("'R' for revenue and 'E' for expenses. If none, enter [n]:")
        if capital_entry == 'R':
            revenue_record()
        elif capital_entry == 'E':
            expense_record()
        elif capital_entry == 'n':
            recording(user_choice)
        else:
            print('That is an invalid syntax. Try Again.')
            start_program()
    else:
        recording(user_choice)

#asks for specific information regarding the user_choice part
def choice_handle(user_choice):
	recording(user_choice)

def Withdrawal_record(): #Recording of withdrawals
    n = int(input("How many Withdrawal entries incurred?: "))

    print("\n")
    for i in range(0, n):
        print("Withdrawal #", i+1, ':', sep='')
        withdraw = input()
        Withdrawal.append(withdraw)
        print("You have entered ", Withdrawal)
        print("Value of Withdrawal #", i+1, ':', sep='')
        withdrawn = int(input())
        Withdrawal_value.append(withdrawn)
    #shows what user have entered
    print('You have entered the following account title(s):',Withdrawal)
    print('You have entered the following asset value:',Withdrawal_value)
    time.sleep(1)
    #user confirmation
    user_confirmation=input('is the following information correct?:')
    if user_confirmation=='y':
        Retry_parameter()
    elif user_confirmation=='n':
        time.sleep(0.5)
        recording(user_choice)    

def expense_record(): #recording of expenses
    n = int(input("How many Expense Entries incurred?: "))

    print("\n")
    for i in range(0, n):
        print("Expense #", i+1, ':', sep='')
        entry = input()
        Expenditures.append(entry)
        print("You have entered ", Expenditures)
        print("Value of Expense #", i+1, ':', sep='')
        value = int(input())
        Expenditure_Cost.append(value)
    #shows what user have entered
    print('You have entered the following account title(s):',Expenditures)
    print('You have entered the following asset value:',Expenditure_Cost)
    time.sleep(1)
    #user confirmation
    user_confirmation=input('is the following information correct?:')
    if user_confirmation=='y':
        Retry_parameter()
    elif user_confirmation=='n':
        time.sleep(0.5)
        recording(user_choice)   

def revenue_record(): #Recording of Revenue
    n = int(input("How many Revenue Entries incurred?: "))

    print("\n")
    for i in range(0, n):
        print("Revenue #", i+1, ':', sep='')
        entry = input()
        Revenues.append(entry)
        print("You have entered ", Revenues)
        print("Value of Revenue #", i+1, ':', sep='')
        value = int(input())
        Revenue_Quantity.append(value)
    #shows what user have entered
    print('You have entered the following account title(s):',Revenues)
    print('You have entered the following asset value:',Revenue_Quantity)
    time.sleep(1)
    #user confirmation
    user_confirmation=input('is the following information correct?:')
    if user_confirmation=='y':
        Retry_parameter()
    elif user_confirmation=='n':
        time.sleep(0.5)
        recording(user_choice)

def intangible_asset_record(): #Recording of Intangible Assets such as goodwill or customer care service ratings.

    n = int(input("How many Assets?: "))

    print("\n")
    for i in range(0, n):
        print("Asset #", i+1, ':', sep='')
        entry = input()
        Intangible_Assets.append(entry)
        print("You have entered ", Intangible_Assets)
        print("Value of Asset #", i+1, ':', sep='')
        value = int(input())
        Intangible_Assets_Value.append(value)
    #shows what user have entered
    print('You have entered the following account title(s):',Intangible_Assets)
    print('You have entered the following asset value:',Intangible_Assets_Value)
    time.sleep(1)
    #user confirmation
    user_confirmation=input('is the following information correct?:')
    if user_confirmation=='y':
        Retry_parameter()
    elif user_confirmation=='n':
        time.sleep(0.5)
        recording(user_choice)

def current_asset_record(): #Recording of current assets such as cash, furnitures, etc.
    n = int(input("How many assets?: "))

    print("\n")
    for i in range(0, n):
        print("Asset #", i+1, ':', sep='')
        entry = input()
        Current_Assets.append(entry)
        print("You have entered ", Current_Assets)
        print("Value of Asset #", i+1, ':', sep='')
        value = int(input())
        Current_Assets_Value.append(value)
 
    #shows what user have entered
    print('You have entered the following account title(s):',Current_Assets)
    print('You have entered the following asset value:',Current_Assets_Value)
    time.sleep(1)
    #user confirmation
    user_confirmation = input('is the following information correct?[y,n]:')
    if user_confirmation == 'y':
        Retry_parameter()
    elif user_confirmation == 'n':
        time.sleep(0.5)
        recording(user_choice)
        
def fixed_asset_record(): #Recording of fixed assets such as building and land.

    n = int(input("How many assets?: "))

    print("\n")
    for i in range(0, n):
        print("Asset #", i+1, ':', sep='')
        entry = input()
        Fixed_Assets.append(entry)
        print("You have entered ", Fixed_Assets)
        print("Value of Asset #", i+1, ':', sep='')
        value = int(input())
        Fixed_Assets_Value.append(value)
    #shows what user have entered
    print('You have entered the following account title(s):',Fixed_Assets)
    print('You have entered the following asset value:',Fixed_Assets_Value)
    time.sleep(1)
    #user confirmation
    user_confirmation=input('is the following information correct?:')
    if user_confirmation=='y':
        Retry_parameter()
    elif user_confirmation=='n':
        time.sleep(0.5)
        recording(user_choice)

def recording(user_choice):
    if user_choice=='Liabilities': #Recording of liability such as accounts and tax payables.
        n = int(input("How many liabilities?: "))

        print("\n")
        for i in range(0, n):
            print("Liability #", i+1, ':', sep='')
            entry = input()
            Current_Liabilities.append(entry)
            print("You have entered ", Current_Liabilities)
            print("Value of Liability #", i+1, ':', sep='')
            value = int(input())
            Current_Liabilities_Amount.append(value)

        #shows what user have entered
        print('You have entered the following account title(s):',Current_Liabilities)
        print('You have entered the following liability values(s)',Current_Liabilities_Amount)
        time.sleep(1)
        #user confirmation
        user_confirmation=input('is the following information correct?[y,n]:')
        if user_confirmation=='y':
            Retry_parameter()
        elif user_confirmation=='n':
            time.sleep(0.5)
            recording(user_choice)
    elif user_choice=='Capital': #Recording of capital/equity of the owner or a shareholder.
        n = int(input("How many assets?: "))

        print("\n")
        for i in range(0, n):
            print("Capital/Equity #", i+1, ':', sep='')
            entry = input()
            Capital_Entry.append(entry)
            print("You have entered ", Capital_Entry)
            print("Value of Capital/Equity #", i+1, ':', sep='')
            value = int(input())
            Capital_Value.append(value)
        #shows what user have entered
        print('You have entered the following account title(s):',Capital_Entry)
        print('You have entered the following capital values(s)',Capital_Value)
        time.sleep(0.5)
        #user confirmation
        user_confirmation=input('is the following information correct?[y,n]:')
        if user_confirmation=='y':
            Retry_parameter()
        elif user_confirmation=='n':
            time.sleep(0.5)
            recording(user_choice)
    else: #Catches all uneccessary inputs to avoid encountering errors.
        print("Only 'l' or 'c' . Try again")
        time.sleep(0.5)
        start_program()

def Retry_parameter(): #Only appears when the recorded values are correct and asks if the user want to record again.
    time.sleep(0.5)
    again=input('Would you like to record again? [y,n]: ')
    if again=='y':
        start_program()
    elif again=='n':
        print('The values you have recorded are now stored on RAM. To avoid loss, do not quit the application.')
        print('-'*20)
        for a, b in zip(Current_Assets, Current_Assets_Value):
            print(a, b)  
        print('-'*20)
        for a, b in zip(Fixed_Assets, Fixed_Assets_Value):
            print(a, b)
        print('-'*20)
        for a, b in zip(Intangible_Assets, Intangible_Assets_Value):
            print(a, b)
        print('-'*20)
        for a, b in zip(Current_Liabilities, Current_Liabilities_Amount):
            print(a, b)
        print('-'*20)
        for a, b in zip(Capital_Entry, Capital_Value):
            print(a, b)
        print('-'*20)
        for a, b in zip(Revenues, Revenue_Quantity):
            print(a,b)
        print('-'*20)
        for a, b in zip(Expenditures, Expenditure_Cost):
            print(a, b)
        print('-'*20)
        for a, b in zip(Withdrawal, Withdrawal_value):
            print(a, b)

#Phase 2 of The Project
def phase2():
    document = Document()
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    section = document.sections[0]

    sectPR =  section._sectPr
    cols = sectPR.xpath('./w:cols')[0]

    cols.set(qn('w:num'), '1')

    document.add_heading('Balance Sheet', 0)

    global content
    content = 'Hello, thank you for using Py-accounting software. You May Delete this Message.'

    def writedocx(content, font_name = 'Times New Roman', font_size = 12, font_bold = False, font_italic = False, font_underline = False, color = RGBColor(0, 0, 0),
                before_spacing = 5, after_spacing = 5, line_spacing = 1.5, keep_together = True, keep_with_next = False, page_break_before = False,
                widow_control = False, align = 'left', style = ''):
        paragraph = document.add_paragraph(str(content))
        paragraph.style = document.styles.add_style(style, WD_STYLE_TYPE.PARAGRAPH)
        font = paragraph.style.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = font_bold
        font.italic = font_italic
        font.underline = font_underline
        font.color.rgb = color
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(before_spacing)
        paragraph_format.space_after = Pt(after_spacing)
        paragraph.line_spacing = line_spacing
        paragraph_format.keep_together = keep_together
        paragraph_format.keep_with_next = keep_with_next
        paragraph_format.page_break_before = page_break_before
        paragraph_format.widow_control = widow_control
        if align.lower() == 'left':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif align.lower() == 'center':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif align.lower() == 'right':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        elif align.lower() == 'justify':
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        else:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    writedocx(content) 

    document.save('Word.docx')

variables()
start_program()

check=input('Would you like to check for imbalances in both credit side and debit side? [y,n]:')
if check=='y':
    if sum(Current_Assets_Value)+sum(Fixed_Assets_Value)+sum(Intangible_Assets_Value)==sum(Current_Liabilities_Amount)+sum(Capital_Value)+sum(Revenue_Quantity)-sum(Expenditure_Cost)-sum(Withdrawal_value):
        time.sleep(0.5)
        print('OK')
    else:
        time.sleep(0.5)
        print('The inputted values are not equal.')
        if sum(Current_Assets_Value)+sum(Fixed_Assets_Value)+sum(Intangible_Assets_Value)>sum(Current_Liabilities_Amount)+sum(Capital_Value)+sum(Revenue_Quantity)-sum(Expenditure_Cost)-sum(Withdrawal_value):
            print('The debit side is higher then the credit side.')
            time.sleep(1)
            print('The debit side is', sum(Current_Assets_Value)+sum(Fixed_Assets_Value)+sum(Intangible_Assets_Value),'and the credit side is', sum(Current_Liabilities_Amount)+sum(Capital_Value)+sum(Revenue_Quantity)-sum(Expenditure_Cost)-sum(Withdrawal_value))
            seeif=input('Would you like to fix the errors? [y,n]:')
            if seeif=='y':
                start_program()
            elif seeif=='n':
                print('ok.')
        if sum(Current_Assets_Value)+sum(Fixed_Assets_Value)+sum(Intangible_Assets_Value)<sum(Current_Liabilities_Amount)+sum(Capital_Value)+sum(Revenue_Quantity)-sum(Expenditure_Cost)-sum(Withdrawal_value):
            print('The credit side is higher than the debit side.')
            time.sleep(1)
            print('The debit side is', sum(Current_Assets_Value)+sum(Fixed_Assets_Value)+sum(Intangible_Assets_Value),'and the credit side is', sum(Current_Liabilities_Amount)+sum(Capital_Value)+sum(Revenue_Quantity)-sum(Expenditure_Cost)-sum(Withdrawal_value))
            seeif=input('Would you like to fix the errors? [y,n]:')
            if seeif=='y':
                start_program()
            elif seeif=='n':
                print('ok.')
else:
    phase2()


